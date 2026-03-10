#!/usr/bin/env python3
"""
Apply DOCX creation/append plan using python-docx.

Supports creating new documents from scratch with tables, headings, paragraphs, and styles.
"""
from __future__ import annotations
import logging

from collections import Counter
from copy import copy
from pathlib import Path
from typing import Any, Dict, List


logger = logging.getLogger(__name__)

_TC_PR_ORDER = [
    "w:tcW",
    "w:gridSpan",
    "w:hMerge",
    "w:vMerge",
    "w:tcBorders",
    "w:shd",
    "w:noWrap",
    "w:tcMar",
    "w:textDirection",
    "w:tcFitText",
    "w:vAlign",
    "w:hideMark",
    "w:headers",
    "w:cellIns",
    "w:cellDel",
    "w:cellMerge",
    "w:tcPrChange",
]

_TC_PR_ORDER_QN: list[str] | None = None

def _tc_pr_order_qn():
    global _TC_PR_ORDER_QN
    if _TC_PR_ORDER_QN is None:
        from docx.oxml.ns import qn
        _TC_PR_ORDER_QN = [qn(tag) for tag in _TC_PR_ORDER]
    return _TC_PR_ORDER_QN


def _insert_tcpr_child(tc_pr, element, tag: str):
    from docx.oxml.ns import qn
    target = qn(tag)
    # Reposition existing element if present.
    for child in list(tc_pr):
        if child.tag == target:
            tc_pr.remove(child)
            element = child
            break

    order = _tc_pr_order_qn()
    if target not in order:
        tc_pr.append(element)
        return element

    target_idx = order.index(target)
    for idx, child in enumerate(list(tc_pr)):
        if child.tag in order and order.index(child.tag) > target_idx:
            tc_pr.insert(idx, element)
            return element

    tc_pr.append(element)
    return element

def _load_docx():
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        return Document, Inches, Pt, Cm, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, qn, OxmlElement
    except Exception as e:
        raise SystemExit(f"python-docx is required for document creation, but it is not available: {e}") from e


def _parse_alignment(align_str: str | None):
    _, _, _, _, WD_ALIGN_PARAGRAPH, _, _, _ = _load_docx()
    if not align_str:
        return None
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(align_str.lower())


def _set_cell_border(cell, **kwargs):
    """Set cell border. kwargs: top, bottom, left, right with values like {"sz": 12, "color": "000000", "val": "single"}"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
    _insert_tcpr_child(tcPr, tcBorders, "w:tcBorders")
    
    for edge in ["top", "left", "bottom", "right"]:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            element.set(qn('w:val'), edge_data.get("val", "single"))
            element.set(qn('w:sz'), str(edge_data.get("sz", 4)))
            element.set(qn('w:color'), edge_data.get("color", "000000"))


def _set_cell_shading(cell, color: str):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.find(qn('w:shd'))
    if shading is None:
        shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)
    _insert_tcpr_child(tcPr, shading, "w:shd")


def _set_chinese_font(run):
    """Set Chinese font for a run."""
    from docx.oxml.ns import qn
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def apply_create_plan(in_docx: Path | None, plan: Dict[str, Any], out_docx: Path) -> Dict[str, Any]:
    """
    Apply creation-oriented plan to create or extend a DOCX.
    
    If in_docx is None or doesn't exist, creates a new document.
    """
    Document, Inches, Pt, Cm, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, qn, OxmlElement = _load_docx()
    
    # Load existing or create new
    if in_docx and in_docx.exists():
        doc = Document(str(in_docx))
    else:
        doc = Document()
        # Set default styles to Chinese font if possible
        try:
            style = doc.styles['Normal']
            style.font.name = 'Microsoft YaHei'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        except Exception:
            logger.warning("Suppressed exception in apply_create_plan.", exc_info=True)
    
    change_log: List[Dict[str, Any]] = []
    
    for op in plan.get("operations", []):
        kind = op.get("op")
        
        if kind == "add_heading":
            text = op.get("text", "")
            level = int(op.get("level", 1))
            alignment = _parse_alignment(op.get("alignment"))
            
            heading = doc.add_heading(level=level)
            run = heading.add_run(text)
            _set_chinese_font(run)
            
            if alignment is not None:
                heading.alignment = alignment
            
            change_log.append({"op": kind, "level": level, "text": text[:50]})

        elif kind == "add_paragraph":
            text = op.get("text", "")
            style = op.get("style")
            alignment = _parse_alignment(op.get("alignment"))
            bold = op.get("bold", False)
            
            para = doc.add_paragraph(style=style)
            run = para.add_run(text)
            _set_chinese_font(run)
            
            if bold:
                run.bold = True
            if alignment is not None:
                para.alignment = alignment
            
            change_log.append({"op": kind, "text": text[:50]})

        elif kind == "add_line_break":
            # Add empty paragraph as line break
            doc.add_paragraph()
            change_log.append({"op": kind})

        elif kind == "create_table":
            rows = int(op.get("rows", 1))
            cols = int(op.get("cols", 1))
            data = op.get("data", [])  # List of rows, each row is list of cell texts
            header_row = op.get("header_row", False)
            style = op.get("style", "Table Grid")
            col_widths = op.get("col_widths", [])  # List of widths in cm
            
            table = doc.add_table(rows=rows, cols=cols)
            try:
                table.style = style
            except Exception:
                logger.warning("Suppressed exception in apply_create_plan.", exc_info=True)
            
            # Set column widths
            if col_widths:
                for i, width in enumerate(col_widths):
                    if i < len(table.columns):
                        for cell in table.columns[i].cells:
                            cell.width = Cm(width)
            
            # Fill data
            for row_idx, row_data in enumerate(data):
                if row_idx >= len(table.rows):
                    break
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx >= len(table.rows[row_idx].cells):
                        break
                    cell = table.rows[row_idx].cells[col_idx]
                    # Direct text set doesn't allow font control easily, so we access the p
                    cell.text = "" # Clear default
                    para = cell.paragraphs[0]
                    run = para.add_run(str(cell_text) if cell_text else "")
                    _set_chinese_font(run)
                    
                    # Header row styling
                    if header_row and row_idx == 0:
                        run.bold = True
            
            change_log.append({"op": kind, "rows": rows, "cols": cols})

        elif kind == "style_table_cell":
            # Style a specific cell in the last created table (by row/col index)
            table_idx = int(op.get("table", -1))  # -1 means last table
            row = int(op.get("row", 0))
            col = int(op.get("col", 0))
            bg_color = op.get("bg_color")
            bold = op.get("bold", False)
            alignment = _parse_alignment(op.get("alignment"))
            
            tables = doc.tables
            if tables:
                tbl = tables[table_idx]
                if row < len(tbl.rows) and col < len(tbl.rows[row].cells):
                    cell = tbl.rows[row].cells[col]
                    if bg_color:
                        _set_cell_shading(cell, bg_color)
                    if bold:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True
                    if alignment is not None:
                        for para in cell.paragraphs:
                            para.alignment = alignment
            
            change_log.append({"op": kind, "row": row, "col": col})

        elif kind == "merge_table_cells":
            # Merge cells in the last table
            table_idx = int(op.get("table", -1))
            start_row = int(op.get("start_row", 0))
            start_col = int(op.get("start_col", 0))
            end_row = int(op.get("end_row", start_row))
            end_col = int(op.get("end_col", start_col))
            
            tables = doc.tables
            if tables:
                tbl = tables[table_idx]
                if start_row < len(tbl.rows) and start_col < len(tbl.rows[start_row].cells):
                    a = tbl.cell(start_row, start_col)
                    b = tbl.cell(end_row, end_col)
                    a.merge(b)
            
            change_log.append({"op": kind, "range": f"({start_row},{start_col})-({end_row},{end_col})"})

        elif kind == "set_table_borders":
            # Set borders for all cells in the last table
            table_idx = int(op.get("table", -1))
            border_style = op.get("style", "single")
            border_size = int(op.get("size", 4))
            border_color = op.get("color", "000000")
            
            tables = doc.tables
            if tables:
                tbl = tables[table_idx]
                border_def = {
                    "top": {"val": border_style, "sz": border_size, "color": border_color},
                    "bottom": {"val": border_style, "sz": border_size, "color": border_color},
                    "left": {"val": border_style, "sz": border_size, "color": border_color},
                    "right": {"val": border_style, "sz": border_size, "color": border_color},
                }
                for row in tbl.rows:
                    for cell in row.cells:
                        _set_cell_border(cell, **border_def)
            
            change_log.append({"op": kind})

        elif kind == "add_page_break":
            doc.add_page_break()
            change_log.append({"op": kind})

        else:
            # Unknown op - skip or raise
            change_log.append({"op": kind, "error": "unknown_op"})
    
    # Save
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    
    counts = Counter([c.get("op") for c in change_log if isinstance(c, dict) and isinstance(c.get("op"), str)])
    change_summary = ", ".join([f"{k}×{counts[k]}" for k in sorted(counts)]) if counts else ""

    errors: list[dict[str, Any]] = []
    for c in change_log:
        if isinstance(c, dict) and c.get("error"):
            errors.append({"message": f"{c.get('op')}: {c.get('error')}"})

    rep: dict[str, Any] = {
        "engine": "python-docx",
        "touched_parts": ["word/document.xml"],
        "operations": change_log,
        "risk_summary": {
            "has_charts": False,
            "has_pivots": False,
            "has_controls": False,
            "has_macros": False,
            "has_formulas": False,
            "has_external_links": False,
            "risk_level": "low",
        },
        "docx_features": {
            "has_tracked_changes": False,
            "has_comments": False,
        },
        "change_summary": change_summary,
    }
    if errors:
        rep["errors"] = errors
    return rep
