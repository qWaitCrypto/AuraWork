#!/usr/bin/env python3
"""
High-level DOCX operations for office workflows.

Design goals:
- Prefer `python-docx` when available for convenience.
- Provide an OOXML fallback path so the most common operations still work
  without extra dependencies.
- For tracked changes / precise comment ranges, prefer OOXML unpack/edit/validate/pack.
"""
from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document  # type: ignore
except Exception as e:  # pragma: no cover
    Document = None  # type: ignore

import xml.etree.ElementTree as ET

from utilities import NS, sha1_text

def _has_python_docx() -> bool:
    return Document is not None

def _require_python_docx() -> None:
    if not _has_python_docx():
        raise RuntimeError("python-docx is not available in this environment.")

def _read_part_from_docx(docx_path: Path, part_name: str) -> bytes:
    with zipfile.ZipFile(docx_path, "r") as zf:
        return zf.read(part_name)

def _parse_document_xml_bytes(xml_bytes: bytes) -> ET.Element:
    # Register stable prefixes to avoid ns0/ns1 output when re-serializing.
    for prefix, uri in NS.items():
        if prefix == "xml":
            continue
        ET.register_namespace(prefix, uri)
    return ET.fromstring(xml_bytes)

def _serialize_xml_root(root: ET.Element) -> bytes:
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)

def _read_text_of_paragraph(p: ET.Element) -> str:
    texts: list[str] = []
    for t in p.findall(".//w:t", NS):
        if t.text:
            texts.append(t.text)
    return "".join(texts).strip()

def _extract_heading_index_from_root(document_root: ET.Element) -> list[dict[str, str]]:
    out: list[dict[str, str]] = []
    for p in document_root.findall(".//w:p", NS):
        ppr = p.find("w:pPr", NS)
        style = None
        if ppr is not None:
            pstyle = ppr.find("w:pStyle", NS)
            if pstyle is not None:
                style = pstyle.attrib.get(f"{{{NS['w']}}}val") or pstyle.attrib.get("w:val") or pstyle.attrib.get("val")
        if not style:
            continue
        text = _read_text_of_paragraph(p)
        if style.lower().startswith("heading") and text:
            out.append({"style": style, "text": text, "anchor": sha1_text(text[:200])})
    return out

def _analyze_docx_ooxml(docx_path: Path) -> Dict[str, Any]:
    xml_bytes = _read_part_from_docx(docx_path, "word/document.xml")
    root = _parse_document_xml_bytes(xml_bytes)

    paragraphs = len(root.findall(".//w:p", NS))
    tables = len(root.findall(".//w:tbl", NS))

    style_counts: dict[str, int] = {}
    for p in root.findall(".//w:p", NS):
        ppr = p.find("w:pPr", NS)
        style = "Unknown"
        if ppr is not None:
            pstyle = ppr.find("w:pStyle", NS)
            if pstyle is not None:
                style = (
                    pstyle.attrib.get(f"{{{NS['w']}}}val")
                    or pstyle.attrib.get("w:val")
                    or pstyle.attrib.get("val")
                    or "Unknown"
                )
        style_counts[style] = style_counts.get(style, 0) + 1

    return {
        "file": str(docx_path),
        "paragraphs": paragraphs,
        "tables": tables,
        "heading_index": _extract_heading_index_from_root(root),
        "style_counts": dict(sorted(style_counts.items(), key=lambda kv: (-kv[1], kv[0]))),
        "backend": "ooxml",
    }

def _replace_text_ooxml(in_docx: Path, out_docx: Path, find: str, replace: str) -> Dict[str, Any]:
    xml_bytes = _read_part_from_docx(in_docx, "word/document.xml")
    root = _parse_document_xml_bytes(xml_bytes)

    replacements = 0
    for t in root.findall(".//w:t", NS):
        if not t.text:
            continue
        if find not in t.text:
            continue
        replacements += t.text.count(find)
        t.text = t.text.replace(find, replace)

    new_document_xml = _serialize_xml_root(root)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(in_docx, "r") as zin, zipfile.ZipFile(out_docx, "w") as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename == "word/document.xml":
                data = new_document_xml
            zout.writestr(info, data)

    return {"replacements": replacements, "out": str(out_docx), "backend": "ooxml"}

def _extract_tables_ooxml(docx_path: Path) -> Dict[str, Any]:
    xml_bytes = _read_part_from_docx(docx_path, "word/document.xml")
    root = _parse_document_xml_bytes(xml_bytes)

    tables: list[list[list[str]]] = []
    for tbl in root.findall(".//w:tbl", NS):
        rows: list[list[str]] = []
        for tr in tbl.findall("./w:tr", NS):
            row: list[str] = []
            for tc in tr.findall("./w:tc", NS):
                cell_paras: list[str] = []
                for p in tc.findall(".//w:p", NS):
                    text = _read_text_of_paragraph(p)
                    if text:
                        cell_paras.append(text)
                row.append("\n".join(cell_paras).strip())
            rows.append(row)
        tables.append(rows)

    return {"file": str(docx_path), "tables": tables, "table_count": len(tables), "backend": "ooxml"}

def _extract_text_ooxml(docx_path: Path) -> Dict[str, Any]:
    xml_bytes = _read_part_from_docx(docx_path, "word/document.xml")
    root = _parse_document_xml_bytes(xml_bytes)

    paras: list[str] = []
    for p in root.findall(".//w:p", NS):
        text = _read_text_of_paragraph(p)
        if text:
            paras.append(text)

    text = "\n".join(paras).strip() + ("\n" if paras else "")
    return {"file": str(docx_path), "paragraphs": len(paras), "text": text, "backend": "ooxml"}

def analyze_docx(path: Path) -> Dict[str, Any]:
    if _has_python_docx():
        try:
            doc = Document(str(path))
            headings: List[Dict[str, str]] = []
            para_count = 0
            table_count = len(doc.tables)

            style_counts: Dict[str, int] = {}

            for p in doc.paragraphs:
                para_count += 1
                style = (p.style.name if p.style is not None else "Unknown")  # type: ignore
                style_counts[style] = style_counts.get(style, 0) + 1
                if style.lower().startswith("heading") and p.text.strip():
                    headings.append({
                        "style": style,
                        "text": p.text.strip(),
                        "anchor": sha1_text(p.text.strip()[:200]),
                    })

            return {
                "file": str(path),
                "paragraphs": para_count,
                "tables": table_count,
                "heading_index": headings,
                "style_counts": dict(sorted(style_counts.items(), key=lambda kv: (-kv[1], kv[0]))),
                "backend": "python-docx",
            }
        except Exception as e:
            res = _analyze_docx_ooxml(path)
            res["note"] = f"python-docx failed; used OOXML fallback: {e}"
            return res

    return _analyze_docx_ooxml(path)

def replace_text(path_in: Path, path_out: Path, find: str, replace: str) -> Dict[str, Any]:
    if find == "":
        raise ValueError("--find must be a non-empty string")

    if _has_python_docx():
        try:
            def _replace_in_runs(runs: Any) -> int:
                hits = 0
                for run in runs:
                    text = getattr(run, "text", "")
                    if not text or find not in text:
                        continue
                    hits += text.count(find)
                    run.text = text.replace(find, replace)
                return hits

            doc = Document(str(path_in))
            hits = 0
            for p in doc.paragraphs:
                hits += _replace_in_runs(p.runs)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            hits += _replace_in_runs(p.runs)
            doc.save(str(path_out))
            return {"replacements": hits, "out": str(path_out), "backend": "python-docx"}
        except Exception as e:
            res = _replace_text_ooxml(path_in, path_out, find, replace)
            res["note"] = f"python-docx failed; used OOXML fallback: {e}"
            return res

    return _replace_text_ooxml(path_in, path_out, find, replace)

def generate_from_outline(outline_path: Path, out_docx: Path, title: Optional[str] = None) -> Dict[str, Any]:
    """
    Generate a DOCX from a simple markdown-like outline:
    - Lines starting with '#', '##', '###' become headings.
    - Other non-empty lines become paragraphs.
    """
    def _write_minimal_docx(*, document_xml: bytes, out_path: Path) -> None:
        content_types = (
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">\n'
            '  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>\n'
            '  <Default Extension="xml" ContentType="application/xml"/>\n'
            '  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>\n'
            "</Types>\n"
        ).encode("utf-8")

        root_rels = (
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">\n'
            '  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>\n'
            "</Relationships>\n"
        ).encode("utf-8")

        # Keep an empty document rels part so Gate A validates it when present.
        document_rels = (
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>\n'
        ).encode("utf-8")

        out_path.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("[Content_Types].xml", content_types)
            zf.writestr("_rels/.rels", root_rels)
            zf.writestr("word/document.xml", document_xml)
            zf.writestr("word/_rels/document.xml.rels", document_rels)

    def _build_document_xml(*, outline_text: str, title_text: Optional[str]) -> bytes:
        for prefix, uri in NS.items():
            if prefix == "xml":
                continue
            ET.register_namespace(prefix, uri)

        def _add_paragraph(body: ET.Element, *, text: str, style: str | None = None) -> None:
            p = ET.SubElement(body, f"{{{NS['w']}}}p")
            if style:
                ppr = ET.SubElement(p, f"{{{NS['w']}}}pPr")
                pstyle = ET.SubElement(ppr, f"{{{NS['w']}}}pStyle")
                pstyle.set(f"{{{NS['w']}}}val", style)
            r = ET.SubElement(p, f"{{{NS['w']}}}r")
            t = ET.SubElement(r, f"{{{NS['w']}}}t")
            if text and (text[0].isspace() or text[-1].isspace()):
                t.set(f"{{{NS['xml']}}}space", "preserve")
            t.text = text

        doc = ET.Element(f"{{{NS['w']}}}document")
        body = ET.SubElement(doc, f"{{{NS['w']}}}body")

        if title_text:
            _add_paragraph(body, text=title_text.strip(), style="Title")

        for raw_line in outline_text.splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                continue
            if line.startswith("### "):
                _add_paragraph(body, text=line[4:].strip(), style="Heading3")
            elif line.startswith("## "):
                _add_paragraph(body, text=line[3:].strip(), style="Heading2")
            elif line.startswith("# "):
                _add_paragraph(body, text=line[2:].strip(), style="Heading1")
            else:
                _add_paragraph(body, text=line.strip())

        # Minimal section properties so document.xml passes schema validation.
        sectPr = ET.SubElement(body, f"{{{NS['w']}}}sectPr")
        pgSz = ET.SubElement(sectPr, f"{{{NS['w']}}}pgSz")
        pgSz.set(f"{{{NS['w']}}}w", "12240")
        pgSz.set(f"{{{NS['w']}}}h", "15840")
        pgMar = ET.SubElement(sectPr, f"{{{NS['w']}}}pgMar")
        pgMar.set(f"{{{NS['w']}}}top", "1440")
        pgMar.set(f"{{{NS['w']}}}right", "1440")
        pgMar.set(f"{{{NS['w']}}}bottom", "1440")
        pgMar.set(f"{{{NS['w']}}}left", "1440")
        pgMar.set(f"{{{NS['w']}}}header", "720")
        pgMar.set(f"{{{NS['w']}}}footer", "720")
        pgMar.set(f"{{{NS['w']}}}gutter", "0")
        cols = ET.SubElement(sectPr, f"{{{NS['w']}}}cols")
        cols.set(f"{{{NS['w']}}}space", "720")
        docGrid = ET.SubElement(sectPr, f"{{{NS['w']}}}docGrid")
        docGrid.set(f"{{{NS['w']}}}linePitch", "360")

        return ET.tostring(doc, encoding="utf-8", xml_declaration=True)

    if _has_python_docx():
        doc = Document()
        if title:
            doc.add_heading(title, level=0)
        raw = outline_path.read_text(encoding="utf-8")
        for line in raw.splitlines():
            line = line.rstrip()
            if not line.strip():
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            else:
                doc.add_paragraph(line.strip())
        doc.save(str(out_docx))
        return {"out": str(out_docx), "backend": "python-docx"}

    raw = outline_path.read_text(encoding="utf-8")
    document_xml = _build_document_xml(outline_text=raw, title_text=title)
    _write_minimal_docx(document_xml=document_xml, out_path=out_docx)
    return {"out": str(out_docx), "backend": "ooxml_minimal"}

def extract_tables(path: Path) -> Dict[str, Any]:
    if _has_python_docx():
        try:
            doc = Document(str(path))
            tables: List[List[List[str]]] = []
            for tbl in doc.tables:
                tdata: List[List[str]] = []
                for row in tbl.rows:
                    tdata.append([cell.text for cell in row.cells])
                tables.append(tdata)
            return {"file": str(path), "tables": tables, "table_count": len(tables), "backend": "python-docx"}
        except Exception as e:
            res = _extract_tables_ooxml(path)
            res["note"] = f"python-docx failed; used OOXML fallback: {e}"
            return res

    return _extract_tables_ooxml(path)

def extract_text(path: Path) -> Dict[str, Any]:
    if _has_python_docx():
        try:
            doc = Document(str(path))
            paras: list[str] = []
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    paras.append(p.text.strip())
            text = "\n".join(paras).strip() + ("\n" if paras else "")
            return {"file": str(path), "paragraphs": len(paras), "text": text, "backend": "python-docx"}
        except Exception as e:
            res = _extract_text_ooxml(path)
            res["note"] = f"python-docx failed; used OOXML fallback: {e}"
            return res

    return _extract_text_ooxml(path)

def main() -> None:
    ap = argparse.ArgumentParser(description="High-level office DOCX operations.")
    sub = ap.add_subparsers(dest="cmd", required=True)

    a = sub.add_parser("analyze", help="Analyze structure and styles.")
    a.add_argument("docx", type=Path)
    a.add_argument("--out", type=Path, default=None)

    r = sub.add_parser("replace", help="Replace text in paragraphs and table cells (no tracked changes).")
    r.add_argument("in_docx", type=Path)
    r.add_argument("out_docx", type=Path)
    r.add_argument("--find", required=True)
    r.add_argument("--replace", required=True)

    g = sub.add_parser("generate", help="Generate docx from outline file.")
    g.add_argument("outline", type=Path)
    g.add_argument("out_docx", type=Path)
    g.add_argument("--title", default=None)

    e = sub.add_parser("extract-tables", help="Extract tables as JSON.")
    e.add_argument("docx", type=Path)
    e.add_argument("--out", type=Path, default=None)

    t = sub.add_parser("extract-text", help="Extract plain text (paragraphs) for summarization/search.")
    t.add_argument("docx", type=Path)
    t.add_argument("--out", type=Path, default=None, help="Write extracted text to this path (UTF-8).")

    args = ap.parse_args()

    if args.cmd == "analyze":
        res = analyze_docx(args.docx)
        if args.out:
            args.out.parent.mkdir(parents=True, exist_ok=True)
            args.out.write_text(json.dumps(res, ensure_ascii=False, indent=2), encoding="utf-8")
        else:
            print(json.dumps(res, ensure_ascii=False, indent=2))
        return

    if args.cmd == "replace":
        res = replace_text(args.in_docx, args.out_docx, args.find, args.replace)
    elif args.cmd == "generate":
        res = generate_from_outline(args.outline, args.out_docx, args.title)
    elif args.cmd == "extract-tables":
        res = extract_tables(args.docx)
        if args.out:
            args.out.parent.mkdir(parents=True, exist_ok=True)
            args.out.write_text(json.dumps(res, ensure_ascii=False, indent=2), encoding="utf-8")
        else:
            print(json.dumps(res, ensure_ascii=False, indent=2))
        return

    if args.cmd == "extract-text":
        res = extract_text(args.docx)
        if args.out:
            args.out.parent.mkdir(parents=True, exist_ok=True)
            args.out.write_text(res["text"], encoding="utf-8")
            print(json.dumps({"file": res["file"], "paragraphs": res["paragraphs"], "out": str(args.out), "backend": res.get("backend")}, ensure_ascii=False))
        else:
            print(res["text"], end="")
        return

    if args.cmd in ("replace", "generate"):
        print(json.dumps(res, ensure_ascii=False, indent=2))
    else:
        raise SystemExit(2)

if __name__ == "__main__":
    main()
